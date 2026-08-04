"""
docx_parser.py
Extracts plain text from PDF (.pdf) or Word (.docx) resume files.
Returns a clean string ready for ATS analysis.
"""

import io
import logging

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: io.BytesIO) -> str:
    """Extract text from a PDF file using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_bytes)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise ValueError(f"Could not read PDF file. It may be corrupted or password-protected. ({e})")


def extract_text_from_docx(file_bytes: io.BytesIO) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(file_bytes)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Could not read DOCX file. It may be corrupted or not a valid Word document. ({e})")


def extract_text(file_bytes: io.BytesIO, filename: str) -> str:
    """
    Route to the correct extractor based on file extension.
    Returns extracted plain text or raises ValueError on failure.
    """
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Unsupported file type. Please upload a .pdf or .docx file.")

    if not text or len(text.strip()) < 50:
        raise ValueError(
            "The file appears to be empty or contains no readable text. "
            "Please ensure the document is not a scanned image without OCR."
        )

    return text.strip()
